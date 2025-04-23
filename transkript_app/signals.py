from django.db.models.signals import post_save
from django.dispatch import receiver

import logging
import os
import subprocess
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH  # Alignment uchun import

import tempfile
from django.core.files import File

from docx.shared import Pt, Inches
from docx.oxml.ns import qn

# Env
from dotenv import load_dotenv

# loading env
load_dotenv()

# QR Code
import qrcode
from io import BytesIO

# Random
import random

# Logging sozlamalari
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

from .models import Transkript


@receiver(post_save, sender=Transkript)
def generate_transkript_pdf(sender, instance, created, **kwargs):
    try:
        doc = Document(instance.yonalish.shablon_docx.path)

        logging.info("Transkript PDF yaratish jarayoni boshlandi.")

        # QR Code
        base_url = "http://192.168.2.109:8000" if os.getenv("DJANGO_ENV") == "dev" else "https://auezovedu.kz"

        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=2,
            border=3,
        )
        qr.add_data(f"{base_url}/checkdocuments/transcript/{instance.student_id}/")
        qr.make(fit=True)

        qr_img = qr.make_image(fill_color="black", back_color="white")

        # QR kodni saqlash uchun vaqtinchalik fayl
        qr_io = BytesIO()
        qr_img.save(qr_io, format='PNG')
        qr_io.seek(0)

        # QR kodni dokumentga qo'shish
        doc.add_picture(qr_io, width=Inches(0.7))
        site_url = doc.add_paragraph(f"{base_url}/checkdocuments/transcript/{instance.student_id}/")
        site_url.style.font.size = Pt(7)
        site_url.style.font.name = 'Calibri'
        site_url.style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

        # Placeholder'lar va ularning shrift o'lchamlari, alignment va bold
        placeholders_config = {
            '{{full_name}}': {
                'value': instance.toliq_ism,
                'size': 7,
                'alignment': None,  # Default alignment
                'bold': True  # Bold format
            },
            '{{faculty}}': {
                'value': instance.fakultet,
                'size': 7,
                'alignment': None,  # Default alignment
                'bold': True  # Bold format
            },
            '{{course_code}}': {
                'value': instance.yonalish.kodi,
                'size': 7,
                'alignment': None,  # Default alignment
                'bold': True  # Bold format
            },
            '{{course_name}}': {
                'value': instance.yonalish.nomi,
                'size': 7,
                'alignment': None,  # Default alignment
                'bold': True  # Bold format
            },
            '{{study_type}}': {
                'value': instance.oqish_turi.nomi,
                'size': 7,
                'alignment': None,  # Default alignment
                'bold': True  # Bold format
            },
            '{{study_course}}': {
                'value': instance.oqish_kursi.nomi,
                'size': 7,
                'alignment': None,  # Default alignment
                'bold': True  # Bold format
            },
            '{{study_lang}}': {
                'value': instance.oqish_tili.nomi,
                'size': 7,
                'alignment': None,  # Default alignment
                'bold': True  # Bold format
            },
            '{{study_finally_year}}': {
                'value': instance.tugatgan_yili,
                'size': 7,
                'alignment': None,  # Default alignment
                'bold': True  # Bold format
            },
            '{{student_id}}': {
                'value': instance.student_id,
                'size': 7,
                'alignment': None,  # Default alignment
                'bold': True  # Bold format
            },
            '{{transcript_id}}': {
                'value': 7400 + instance.id,
                'size': 7,
                'alignment': None,  # Default alignment
                'bold': True  # Bold format
            },
        }

        # Barcha o'zgaruvchilarni yaratish
        # 1-semestr
        point_summa_1 = 0
        ball_summa_1 = 0
        for i in range(1, 13):
            study_credit_data = "{{study_credit_1_" + f"{i}" + "}}"
            study_credit = random.randint(3, 6)  # Kredit soatlarini tasodifiy tanlash
            placeholders_config[study_credit_data] = {
                'value': str(study_credit),
                'size': 4,
                'alignment': WD_ALIGN_PARAGRAPH.CENTER,  # Markaz alignment
                'bold': False  # Bold emas
            }

            ball_data = "{{ball_1_" + f"{i}" + "}}"
            ball = random.randint(56, 95)  # Ballarni tasodifiy tanlash
            placeholders_config[ball_data] = {
                'value': str(ball),
                'size': 4,
                'alignment': WD_ALIGN_PARAGRAPH.CENTER,  # Markaz alignment
                'bold': False  # Bold emas
            }

            # Baholarni ballga qarab tanlash
            if ball >= 90:
                letter = 'C'
            elif ball >= 85:
                letter = 'B+'
            elif ball >= 80:
                letter = 'B'
            elif ball >= 70:
                letter = 'A+'
            else:
                letter = 'A'

            ball_summa_1 += ball

            letter_data = "{{letter_1_" + f"{i}" + "}}"
            placeholders_config[letter_data] = {
                'value': letter,
                'size': 4,
                'alignment': WD_ALIGN_PARAGRAPH.CENTER,  # Markaz alignment
                'bold': False  # Bold emas
            }

            # Pointlar bahoga qarab hisoblanadi

            if letter == 'C+':
                point = round(random.uniform(4.75, 5.0), 2)  # 5 ballik tizimda A+ (4.75 - 5.0)
            elif letter == 'C':
                point = round(random.uniform(4.5, 4.74), 2)  # 5 ballik tizimda A (4.5 - 4.74)
            elif letter == 'B+':
                point = round(random.uniform(4.0, 4.49), 2)  # 5 ballik tizimda B+ (4.0 - 4.49)
            elif letter == 'B':
                point = round(random.uniform(3.5, 3.99), 2)  # 5 ballik tizimda B (3.5 - 3.99)
            elif letter == 'A+':
                point = round(random.uniform(3.0, 3.49), 2)  # 5 ballik tizimda C+ (3.0 - 3.49)
            elif letter == 'A':
                point = round(random.uniform(2.5, 2.99), 2)  # 5 ballik tizimda C (2.5 - 2.99)
            elif letter == 'D+':
                point = round(random.uniform(2.0, 2.49), 2)  # 5 ballik tizimda D+ (2.0 - 2.49)
            else:
                point = round(random.uniform(0.0, 1.99), 2)  # 5 ballik tizimda F (0.0 - 1.99)

            point_summa_1 += point

            point_data = "{{point_1_" + f"{i}" + "}}"
            placeholders_config[point_data] = {
                'value': str(point),
                'size': 4,
                'alignment': WD_ALIGN_PARAGRAPH.CENTER,  # Markaz alignment
                'bold': False  # Bold emas
            }

            # Traditsion baho ballarga qarab tanlanadi
            if point >= 3.5:
                traditional = "5(отлично)"
            elif point >= 3.0:
                traditional = "4(xорошо)"
            else:
                traditional = "3(удов)"

            traditional_data = "{{traditional_1_" + f"{i}" + "}}"
            placeholders_config[traditional_data] = {
                'value': traditional,
                'size': 4,
                'alignment': WD_ALIGN_PARAGRAPH.CENTER,  # Markaz alignment
                'bold': False  # Bold emas
            }

        point_summa_1 = point_summa_1 / 12
        ball_summa_1 = ball_summa_1 / 12

        # 2-semestr
        point_summa_2 = 0
        ball_summa_2 = 0
        for i in range(1, 9):
            study_credit_data = "{{study_credit_2_" + f"{i}" + "}}"
            study_credit = random.randint(3, 6)  # Kredit soatlarini tasodifiy tanlash
            placeholders_config[study_credit_data] = {
                'value': str(study_credit),
                'size': 4,
                'alignment': WD_ALIGN_PARAGRAPH.CENTER,  # Markaz alignment
                'bold': False  # Bold emas
            }

            ball_data = "{{ball_2_" + f"{i}" + "}}"
            ball = random.randint(56, 95)  # Ballarni tasodifiy tanlash
            placeholders_config[ball_data] = {
                'value': str(ball),
                'size': 4,
                'alignment': WD_ALIGN_PARAGRAPH.CENTER,  # Markaz alignment
                'bold': False  # Bold emas
            }

            # Baholarni ballga qarab tanlash
            if ball >= 90:
                letter = 'C'
            elif ball >= 85:
                letter = 'B+'
            elif ball >= 80:
                letter = 'B'
            elif ball >= 70:
                letter = 'A+'
            else:
                letter = 'A'

            ball_summa_2 += ball

            letter_data = "{{letter_2_" + f"{i}" + "}}"
            placeholders_config[letter_data] = {
                'value': letter,
                'size': 4,
                'alignment': WD_ALIGN_PARAGRAPH.CENTER,  # Markaz alignment
                'bold': False  # Bold emas
            }

            # Pointlar bahoga qarab hisoblanadi
            if letter == 'C+':
                point = round(random.uniform(4.75, 5.0), 2)  # 5 ballik tizimda A+ (4.75 - 5.0)
            elif letter == 'C':
                point = round(random.uniform(4.5, 4.74), 2)  # 5 ballik tizimda A (4.5 - 4.74)
            elif letter == 'B+':
                point = round(random.uniform(4.0, 4.49), 2)  # 5 ballik tizimda B+ (4.0 - 4.49)
            elif letter == 'B':
                point = round(random.uniform(3.5, 3.99), 2)  # 5 ballik tizimda B (3.5 - 3.99)
            elif letter == 'A+':
                point = round(random.uniform(3.0, 3.49), 2)  # 5 ballik tizimda C+ (3.0 - 3.49)
            elif letter == 'A':
                point = round(random.uniform(2.5, 2.99), 2)  # 5 ballik tizimda C (2.5 - 2.99)
            elif letter == 'D+':
                point = round(random.uniform(2.0, 2.49), 2)  # 5 ballik tizimda D+ (2.0 - 2.49)
            else:
                point = round(random.uniform(0.0, 1.99), 2)  # 5 ballik tizimda F (0.0 - 1.99)

            point_summa_2 += point

            point_data = "{{point_2_" + f"{i}" + "}}"
            placeholders_config[point_data] = {
                'value': str(point),
                'size': 4,
                'alignment': WD_ALIGN_PARAGRAPH.CENTER,  # Markaz alignment
                'bold': False  # Bold emas
            }

            # Traditsion baho ballarga qarab tanlanadi
            if point >= 3.5:
                traditional = "5(отлично)"
            elif point >= 3.0:
                traditional = "4(xорошо)"
            else:
                traditional = "3(удов)"

            traditional_data = "{{traditional_2_" + f"{i}" + "}}"
            placeholders_config[traditional_data] = {
                'value': traditional,
                'size': 4,
                'alignment': WD_ALIGN_PARAGRAPH.CENTER,  # Markaz alignment
                'bold': False  # Bold emas
            }

        point_summa_2 = point_summa_2 / 8
        ball_summa_2 = ball_summa_2 / 8

        course_1_gpa = (point_summa_1 + point_summa_2) / 2
        course_1_ball = (ball_summa_1 + ball_summa_2) / 2

        placeholders_config['{{course_1_gpa}}'] = {
            'value': str(round(course_1_gpa, 2)),
            'size': 4,
            'alignment': WD_ALIGN_PARAGRAPH.CENTER,  # Markaz alignment
            'bold': False  # Bold emas
        }
        placeholders_config['{{course_1_ball}}'] = {
            'value': str(int(course_1_ball)),
            'size': 4,
            'alignment': WD_ALIGN_PARAGRAPH.CENTER,  # Markaz alignment
            'bold': False  # Bold emas
        }

        # 3-semestr
        point_summa_3 = 0
        ball_summa_3 = 0
        for i in range(1, 13):
            study_credit_data = "{{study_credit_3_" + f"{i}" + "}}"
            study_credit = random.randint(3, 6)  # Kredit soatlarini tasodifiy tanlash
            placeholders_config[study_credit_data] = {
                'value': str(study_credit),
                'size': 4,
                'alignment': WD_ALIGN_PARAGRAPH.CENTER,  # Markaz alignment
                'bold': False  # Bold emas
            }

            ball_data = "{{ball_3_" + f"{i}" + "}}"
            ball = random.randint(56, 95)  # Ballarni tasodifiy tanlash
            placeholders_config[ball_data] = {
                'value': str(ball),
                'size': 4,
                'alignment': WD_ALIGN_PARAGRAPH.CENTER,  # Markaz alignment
                'bold': False  # Bold emas
            }

            # Baholarni ballga qarab tanlash
            if ball >= 90:
                letter = 'C'
            elif ball >= 85:
                letter = 'B+'
            elif ball >= 80:
                letter = 'B'
            elif ball >= 70:
                letter = 'A+'
            else:
                letter = 'A'

            ball_summa_3 += ball

            letter_data = "{{letter_3_" + f"{i}" + "}}"
            placeholders_config[letter_data] = {
                'value': letter,
                'size': 4,
                'alignment': WD_ALIGN_PARAGRAPH.CENTER,  # Markaz alignment
                'bold': False  # Bold emas
            }

            # Pointlar bahoga qarab hisoblanadi
            if letter == 'C+':
                point = round(random.uniform(4.75, 5.0), 2)  # 5 ballik tizimda A+ (4.75 - 5.0)
            elif letter == 'C':
                point = round(random.uniform(4.5, 4.74), 2)  # 5 ballik tizimda A (4.5 - 4.74)
            elif letter == 'B+':
                point = round(random.uniform(4.0, 4.49), 2)  # 5 ballik tizimda B+ (4.0 - 4.49)
            elif letter == 'B':
                point = round(random.uniform(3.5, 3.99), 2)  # 5 ballik tizimda B (3.5 - 3.99)
            elif letter == 'A+':
                point = round(random.uniform(3.0, 3.49), 2)  # 5 ballik tizimda C+ (3.0 - 3.49)
            elif letter == 'A':
                point = round(random.uniform(2.5, 2.99), 2)  # 5 ballik tizimda C (2.5 - 2.99)
            elif letter == 'D+':
                point = round(random.uniform(2.0, 2.49), 2)  # 5 ballik tizimda D+ (2.0 - 2.49)
            else:
                point = round(random.uniform(0.0, 1.99), 2)  # 5 ballik tizimda F (0.0 - 1.99)

            point_summa_3 += point

            point_data = "{{point_3_" + f"{i}" + "}}"
            placeholders_config[point_data] = {
                'value': str(point),
                'size': 4,
                'alignment': WD_ALIGN_PARAGRAPH.CENTER,  # Markaz alignment
                'bold': False  # Bold emas
            }

            # Traditsion baho ballarga qarab tanlanadi
            if point >= 3.5:
                traditional = "5(отлично)"
            elif point >= 3.0:
                traditional = "4(xорошо)"
            else:
                traditional = "3(удов)"

            traditional_data = "{{traditional_3_" + f"{i}" + "}}"
            placeholders_config[traditional_data] = {
                'value': traditional,
                'size': 4,
                'alignment': WD_ALIGN_PARAGRAPH.CENTER,  # Markaz alignment
                'bold': False  # Bold emas
            }

        point_summa_3 = point_summa_3 / 12
        ball_summa_3 = ball_summa_3 / 12

        # 4-semestr
        point_summa_4 = 0
        ball_summa_4 = 0
        for i in range(1, 9):
            study_credit_data = "{{study_credit_4_" + f"{i}" + "}}"
            study_credit = random.randint(3, 6)  # Kredit soatlarini tasodifiy tanlash
            placeholders_config[study_credit_data] = {
                'value': str(study_credit),
                'size': 4,
                'alignment': WD_ALIGN_PARAGRAPH.CENTER,  # Markaz alignment
                'bold': False  # Bold emas
            }

            ball_data = "{{ball_4_" + f"{i}" + "}}"
            ball = random.randint(56, 95)  # Ballarni tasodifiy tanlash
            placeholders_config[ball_data] = {
                'value': str(ball),
                'size': 4,
                'alignment': WD_ALIGN_PARAGRAPH.CENTER,  # Markaz alignment
                'bold': False  # Bold emas
            }

            # Baholarni ballga qarab tanlash
            if ball >= 90:
                letter = 'C'
            elif ball >= 85:
                letter = 'B+'
            elif ball >= 80:
                letter = 'B'
            elif ball >= 70:
                letter = 'A+'
            else:
                letter = 'A'

            ball_summa_4 += ball

            letter_data = "{{letter_4_" + f"{i}" + "}}"
            placeholders_config[letter_data] = {
                'value': letter,
                'size': 4,
                'alignment': WD_ALIGN_PARAGRAPH.CENTER,  # Markaz alignment
                'bold': False  # Bold emas
            }

            # Pointlar bahoga qarab hisoblanadi
            if letter == 'C+':
                point = round(random.uniform(4.75, 5.0), 2)  # 5 ballik tizimda A+ (4.75 - 5.0)
            elif letter == 'C':
                point = round(random.uniform(4.5, 4.74), 2)  # 5 ballik tizimda A (4.5 - 4.74)
            elif letter == 'B+':
                point = round(random.uniform(4.0, 4.49), 2)  # 5 ballik tizimda B+ (4.0 - 4.49)
            elif letter == 'B':
                point = round(random.uniform(3.5, 3.99), 2)  # 5 ballik tizimda B (3.5 - 3.99)
            elif letter == 'A+':
                point = round(random.uniform(3.0, 3.49), 2)  # 5 ballik tizimda C+ (3.0 - 3.49)
            elif letter == 'A':
                point = round(random.uniform(2.5, 2.99), 2)  # 5 ballik tizimda C (2.5 - 2.99)
            elif letter == 'D+':
                point = round(random.uniform(2.0, 2.49), 2)  # 5 ballik tizimda D+ (2.0 - 2.49)
            else:
                point = round(random.uniform(0.0, 1.99), 2)  # 5 ballik tizimda F (0.0 - 1.99)

            point_summa_4 += point

            point_data = "{{point_4_" + f"{i}" + "}}"
            placeholders_config[point_data] = {
                'value': str(point),
                'size': 4,
                'alignment': WD_ALIGN_PARAGRAPH.CENTER,  # Markaz alignment
                'bold': False  # Bold emas
            }

            # Traditsion baho ballarga qarab tanlanadi
            if point >= 3.5:
                traditional = "5(отлично)"
            elif point >= 3.0:
                traditional = "4(xорошо)"
            else:
                traditional = "3(удов)"

            traditional_data = "{{traditional_4_" + f"{i}" + "}}"
            placeholders_config[traditional_data] = {
                'value': traditional,
                'size': 4,
                'alignment': WD_ALIGN_PARAGRAPH.CENTER,  # Markaz alignment
                'bold': False  # Bold emas
            }

        point_summa_4 = point_summa_4 / 8
        ball_summa_4 = ball_summa_4 / 8

        course_2_gpa = (point_summa_3 + point_summa_4) / 2
        course_2_ball = (ball_summa_3 + ball_summa_4) / 2

        placeholders_config['{{course_2_gpa}}'] = {
            'value': str(round(course_2_gpa, 2)),
            'size': 4,
            'alignment': WD_ALIGN_PARAGRAPH.CENTER,  # Markaz alignment
            'bold': False  # Bold emas
        }
        placeholders_config['{{course_2_ball}}'] = {
            'value': str(int(course_2_ball)),
            'size': 4,
            'alignment': WD_ALIGN_PARAGRAPH.CENTER,  # Markaz alignment
            'bold': False  # Bold emas
        }

        placeholders_config['{{total_ball}}'] = {
            'value': str(int(course_1_ball + course_2_ball)),
            'size': 7,
            'alignment': False,
            'bold': True  # Bold
        }

        placeholders_config['{{aw_gpa}}'] = {
            'value': str(round(((course_1_gpa + course_2_gpa) / 2), 2)),
            'size': 7,
            'alignment': False,
            'bold': True  # Bold
        }

        placeholders_config['{{total_subjects}}'] = {
            'value': str(12 + 12 + 8 + 8),
            'size': 7,
            'alignment': False,
            'bold': False  # Bold
        }

        # Barcha paragraflarni tekshirish va placeholder'larni almashtirish
        for paragraph in doc.paragraphs:
            replace_placeholders_with_formatting(paragraph, placeholders_config)

        # Barcha jadvallarni tekshirish va placeholder'larni almashtirish
        process_tables_with_formatting(doc.tables, placeholders_config)

        with tempfile.TemporaryDirectory() as temp_dir:
            temp_docx_path = os.path.join(temp_dir, f'transcript_{instance.student_id}.docx')
            doc.save(temp_docx_path)

            temp_pdf_path = os.path.join(temp_dir, f'transcript_{instance.student_id}.pdf')

            try:
                libreoffice_commands = [
                    "libreoffice", "soffice", "openoffice",
                    "/Applications/LibreOffice.app/Contents/MacOS/soffice",  # macOS uchun
                    "C:\\Program Files\\LibreOffice\\program\\soffice.exe",  # Windows uchun
                ]

                success = False
                for cmd in libreoffice_commands:
                    try:
                        result = subprocess.run(
                            [cmd, "--headless", "--convert-to", "pdf", "--outdir", temp_dir, temp_docx_path],
                            stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=30
                        )
                        if result.returncode == 0:
                            success = True
                            pdf_filename = os.path.splitext(os.path.basename(temp_docx_path))[0] + ".pdf"
                            temp_pdf_path = os.path.join(temp_dir, pdf_filename)
                            break
                    except (subprocess.SubprocessError, FileNotFoundError):
                        continue

                if not success:
                    logging.error("LibreOffice/OpenOffice dasturini topa olmadik yoki ishga tushira olmadik")
                    raise Exception("LibreOffice/OpenOffice dasturini topa olmadik yoki ishga tushira olmadik")

                if os.path.exists(temp_pdf_path):
                    with open(temp_pdf_path, 'rb') as pdf_file:
                        instance.transkript_pdf.save(
                            f'transkript_{instance.student_id}.pdf',
                            File(pdf_file),
                            save=False
                        )

                        Transkript.objects.filter(id=instance.id).update(
                            transkript_pdf=instance.transkript_pdf.name
                        )
                else:
                    logging.error(f"PDF fayl yaratilmadi: {temp_pdf_path}")
                    raise FileNotFoundError(f"PDF fayl yaratilmadi: {temp_pdf_path}")

            except Exception as conversion_error:
                logging.error(f"PDF konvertatsiyasida xatolik: {conversion_error}")
                print(f"PDF konvertatsiyasida xatolik: {conversion_error}")
                raise

    except Exception as e:
        logging.error(f"Transkript PDF yaratishda xatolik: {e}")
        print(f"Transkript PDF yaratishda xatolik: {e}")
        import traceback
        traceback.print_exc()


def replace_placeholders_with_formatting(paragraph, placeholders_config):
    """
    Paragraf ichidagi har bir placeholder uchun to'liq formatlash (o'lcham, markazlashtirish, bold) bilan almashtirish
    """
    if not any(placeholder in paragraph.text for placeholder in placeholders_config.keys()):
        return

    original_text = paragraph.text
    any_replaced = False

    # Har bir placeholder ni tekshirish
    for placeholder, config in placeholders_config.items():
        if placeholder in original_text:
            # Placeholder textini almashtirish
            new_text = original_text.replace(placeholder, str(config['value']))

            # Paragraf tarkibini tozalash
            p_obj = paragraph._p
            for child in list(p_obj):
                p_obj.remove(child)

            # Yangi run qo'shish va formatni o'rnatish
            run = paragraph.add_run(new_text)
            run.font.name = 'Calibri'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
            run.font.size = Pt(config['size'])

            # Bold formati (agar ko'rsatilgan bo'lsa)
            if config.get('bold'):
                run.bold = True

            # Alignment qo'yish (agar ko'rsatilgan bo'lsa)
            if config.get('alignment') is not None:
                paragraph.alignment = config['alignment']

            any_replaced = True
            break  # Bir placeholder almashtirilgandan keyin chiqamiz

    # Agar hech qanday placeholder almashtirilmagan bo'lsa, original textni qaytaramiz
    if not any_replaced:
        p_obj = paragraph._p
        for child in list(p_obj):
            p_obj.remove(child)

        run = paragraph.add_run(original_text)
        run.font.name = 'Calibri'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')


def process_tables_with_formatting(tables, placeholders_config):
    """
    Barcha jadval va ularning ichidagi jadvallarda placeholder'larni almashtirish
    """
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                # Ichki jadvallarda placeholder'larni almashtirish
                # Cell ichidagi jadvallarni tekshirish
                nested_tables = find_nested_tables(cell)
                if nested_tables:
                    process_tables_with_formatting(nested_tables, placeholders_config)

                # Cell ichidagi paragraflarni almashtirish
                for paragraph in cell.paragraphs:
                    replace_placeholders_with_formatting(paragraph, placeholders_config)


def find_nested_tables(cell):
    """
    Cell ichidagi barcha ichki jadvallarni topish
    """
    nested_tables = []
    # Agar cell._element attributi mavjud bo'lsa
    if hasattr(cell, '_element'):
        # Cell ichidagi barcha table elementlarini topish
        for tbl in cell._element.xpath('.//w:tbl'):
            from docx.table import Table
            nested_tables.append(Table(tbl, cell._parent))
    return nested_tables


def remove_table(doc, table_index):
    """
    Word faylidagi berilgan indeksdagi jadvalni o'chiradi.
    """
    tables = doc.tables
    if 0 <= table_index < len(tables):
        tbl = tables[table_index]._element
        tbl.getparent().remove(tbl)